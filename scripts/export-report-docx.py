from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_markdown_inline(paragraph, text: str):
    pos = 0
    for match in LINK_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(clean_inline(text[pos:match.start()]))
        add_hyperlink(paragraph, match.group(1), match.group(2))
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(clean_inline(text[pos:]))


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def add_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [row for index, row in enumerate(rows) if index != 1]
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, cell in enumerate(rows[0]):
        table.rows[0].cells[i].text = clean_inline(cell)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, cell in enumerate(row[: len(cells)]):
            add_markdown_inline(cells[i].paragraphs[0], cell)


def export_docx(markdown_path: Path, output_path: Path, image_path: Path | None = None) -> None:
    markdown = markdown_path.read_text(encoding="utf-8")
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        if line.startswith("# "):
            doc.add_heading(clean_inline(line[2:].strip()), level=0)
        elif line.startswith("## "):
            doc.add_heading(clean_inline(line[3:].strip()), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:].strip()), level=2)
        elif line.startswith("#### "):
            doc.add_heading(clean_inline(line[5:].strip()), level=3)
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_markdown_inline(paragraph, line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_markdown_inline(paragraph, re.sub(r"^\d+\.\s+", "", line).strip())
        elif line.startswith("!["):
            match = IMAGE_RE.match(line)
            if image_path and image_path.exists():
                paragraph = doc.add_paragraph()
                paragraph.add_run().add_picture(str(image_path), width=Pt(430))
                if match and match.group(1):
                    caption = doc.add_paragraph(match.group(1))
                    caption.style = "Caption"
            else:
                paragraph = doc.add_paragraph()
                add_markdown_inline(paragraph, line)
        else:
            paragraph = doc.add_paragraph()
            add_markdown_inline(paragraph, line)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) not in (3, 4):
        raise SystemExit("Usage: export-report-docx.py input.md output.docx [image.png]")
    export_docx(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]) if len(sys.argv) == 4 else None)
